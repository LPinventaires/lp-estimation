import os
import csv
import io
import base64
from datetime import datetime
from functools import wraps
from typing import Optional
from flask import (Flask, render_template, request, redirect, url_for,
                   session, flash, jsonify, send_file)
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from pydantic import BaseModel, Field

from parsing import extract_text, extract_fields
from report import build_report, QUARTIERS

app = Flask(__name__)
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", "dev-secret-change-me")
app.config["MAX_CONTENT_LENGTH"] = 30 * 1024 * 1024

# Upload folder configuration
UPLOAD_FOLDER = os.path.join(app.static_folder, "uploads")
ALLOWED_EXTENSIONS = {"png", "jpg", "jpeg", "gif"}
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

db_url = os.environ.get("DATABASE_URL", "sqlite:///lp_estim.db")
if db_url.startswith("postgres://"):
    db_url = db_url.replace("postgres://", "postgresql://", 1)
app.config["SQLALCHEMY_DATABASE_URI"] = db_url
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
db = SQLAlchemy(app)

# Code d'accès — à changer en prod via APP_PASSWORD
APP_PASSWORD = os.environ.get("APP_PASSWORD", "LP-estimation")


# ----------------------- MODÈLES -----------------------
class RefPrice(db.Model):
    """Référence prix/m² par adresse / quartier (issue des estimations LP et du marché)."""
    id = db.Column(db.Integer, primary_key=True)
    quartier = db.Column(db.String(120))
    adresse = db.Column(db.String(200))
    prix_m2 = db.Column(db.Float)
    annee = db.Column(db.String(20))
    source = db.Column(db.String(200))
    kind = db.Column(db.String(10), default="sold")  # sold | forsale | retenu
    surface = db.Column(db.Float, nullable=True)          # m², si connu
    type_bien = db.Column(db.String(80), nullable=True)   # Appartement, Villa, etc.
    prix_total = db.Column(db.Float, nullable=True)       # prix total CHF, si connu
    description = db.Column(db.Text, nullable=True)       # description riche (comparables enrichis)
    reference = db.Column(db.String(80), nullable=True)   # référence interne (ex: HP12134B)


class Estimation(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), index=True)
    report_ai = db.Column(db.Text, nullable=True)  # narratif LP généré par Claude (mis en cache)
    address = db.Column(db.String(200))
    quartier = db.Column(db.String(120))
    type_bien = db.Column(db.String(80))
    surface = db.Column(db.Float, default=0)
    pieces = db.Column(db.String(30))
    etage = db.Column(db.String(40))
    annee = db.Column(db.String(40))
    etat = db.Column(db.String(160))
    balcon = db.Column(db.Float, default=0)
    balcon_pond = db.Column(db.Float, default=0.5)
    parking_nb = db.Column(db.Integer, default=0)
    parking_val = db.Column(db.Float, default=0)
    prix_m2 = db.Column(db.Float, default=0)
    marge = db.Column(db.Float, default=0.07)
    description = db.Column(db.Text)
    atouts = db.Column(db.Text)
    inconvenients = db.Column(db.Text)
    courtier = db.Column(db.String(120))
    notes = db.Column(db.Text, default="")

    @property
    def val_principale(self):
        return (self.prix_m2 or 0) * (self.surface or 0)

    @property
    def val_balcon(self):
        return (self.prix_m2 or 0) * (self.balcon or 0) * (self.balcon_pond or 0)

    @property
    def val_parking(self):
        return (self.parking_nb or 0) * (self.parking_val or 0)

    @property
    def valeur_venale(self):
        return self.val_principale + self.val_balcon + self.val_parking

    @property
    def prix_presentation(self):
        return round(self.valeur_venale * (1 + (self.marge or 0)) / 10000) * 10000


class Setting(db.Model):
    """Application settings — logo filename, etc."""
    id = db.Column(db.Integer, primary_key=True)
    key = db.Column(db.String(50), unique=True, nullable=False)
    value = db.Column(db.String(500))


class AuditLog(db.Model):
    """Historique des actions pour audit."""
    id = db.Column(db.Integer, primary_key=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    action = db.Column(db.String(100))  # created, updated, deleted, cloned
    estimation_id = db.Column(db.Integer, db.ForeignKey('estimation.id'))
    description = db.Column(db.Text)


class PriceAlert(db.Model):
    """Alertes pour changements de prix par quartier."""
    id = db.Column(db.Integer, primary_key=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    quartier = db.Column(db.String(120))
    previous_price = db.Column(db.Float)
    new_price = db.Column(db.Float)
    triggered = db.Column(db.Boolean, default=False)


class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    email = db.Column(db.String(200), unique=True, nullable=False, index=True)
    password_hash = db.Column(db.String(255), nullable=False)
    name = db.Column(db.String(120))

    def set_password(self, password):
        self.password_hash = generate_password_hash(password, method="pbkdf2:sha256")

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)


# ----------------------- AUTH -----------------------
def login_required(f):
    @wraps(f)
    def wrap(*a, **k):
        if not session.get("user_id"):
            return redirect(url_for("login"))
        return f(*a, **k)
    return wrap


def current_user():
    uid = session.get("user_id")
    if not uid:
        return None
    return User.query.get(uid)


def own_estimations():
    """Query base — estimations de l'utilisateur connecté seulement."""
    return Estimation.query.filter_by(user_id=session.get("user_id"))


def own_estimation_or_404(eid):
    """Fetch une estimation appartenant à l'utilisateur connecté, sinon redirige
    gentiment vers le Classeur (l'estimation a peut-être été supprimée)."""
    e = own_estimations().filter_by(id=eid).first()
    if not e:
        flash("Cette estimation n'existe plus. Voici ton Classeur.")
        from flask import abort
        abort(redirect(url_for("classeur")))
    return e


DEFAULT_LOGO_FILENAME = "leonard-logo.png"


def get_logo_path():
    """Retourne le chemin du logo actuel ou le logo par défaut Leonard Properties."""
    setting = Setting.query.filter_by(key="logo_filename").first()
    if setting and setting.value:
        full_path = os.path.join(app.config["UPLOAD_FOLDER"], setting.value)
        if os.path.exists(full_path):
            return setting.value
    # Fallback : logo par défaut committé dans le repo
    default_full = os.path.join(app.config["UPLOAD_FOLDER"], DEFAULT_LOGO_FILENAME)
    if os.path.exists(default_full):
        return DEFAULT_LOGO_FILENAME
    return None


def allowed_file(filename):
    """Vérifie si le fichier a une extension autorisée."""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def log_audit(action, est_id, description=""):
    """Enregistre une action dans l'audit log."""
    log = AuditLog(action=action, estimation_id=est_id, description=description)
    db.session.add(log)
    db.session.commit()


SHARED_EMAIL = "team@leonardproperties.local"


def _shared_user():
    """Compte partagé Leonard Properties : tout le monde se connecte dessus."""
    u = User.query.filter_by(email=SHARED_EMAIL).first()
    if not u:
        u = User(email=SHARED_EMAIL, name="Leonard Properties")
        u.set_password(os.environ.get("APP_PASSWORD", "LP-estimation"))
        db.session.add(u)
        db.session.commit()
    return u


@app.route("/login", methods=["GET", "POST"])
def login():
    if session.get("user_id"):
        return redirect(url_for("dashboard"))
    expected = (os.environ.get("APP_PASSWORD") or "LP-estimation").strip()
    if request.method == "POST":
        password = (request.form.get("password") or "").strip()
        if password and password == expected:
            session["user_id"] = _shared_user().id
            return redirect(url_for("dashboard"))
        flash("Mot de passe incorrect.")
    return render_template("login.html", logo=get_logo_path())


@app.route("/signup")
def signup():
    """Ancienne route d'inscription — désactivée, redirige vers login."""
    return redirect(url_for("login"))


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("index"))


IMPORT_LP_NOTE = "Importé depuis rapport LP"


@app.route("/admin/import-lp-estimations")
@login_required
def admin_import_lp_estimations():
    """Injecte les estimations LP passées (data/comparables_estimations_lp.csv) dans la table Estimation
    de l'utilisateur connecté, pour qu'elles apparaissent dans le Classeur.
    Idempotent : purge d'abord les précédents imports par leur note distinctive."""
    import csv as _csv, re as _re
    path = os.path.join(os.path.dirname(__file__), "data", "comparables_estimations_lp.csv")
    if not os.path.exists(path):
        return jsonify({"error": "CSV absent"}), 404

    uid = session.get("user_id")
    # Purge des imports précédents (idempotent)
    Estimation.query.filter_by(user_id=uid, notes=IMPORT_LP_NOTE).delete()
    db.session.commit()

    inserted = 0
    skipped = 0
    with open(path, encoding="utf-8") as f:
        for row in _csv.DictReader(f):
            adresse = " ".join((row.get("address") or "").split()).strip()[:200]
            if not adresse:
                skipped += 1
                continue
            type_bien = _normalize_type(row.get("type_bien") or "")[:80]
            if not type_bien:
                skipped += 1
                continue
            surface = _parse_swiss_number(row.get("surface"))
            prix_m2 = _parse_swiss_number(row.get("prix_m2_retenu"))
            valeur = _parse_swiss_number(row.get("valeur_venale"))
            pres = _parse_swiss_number(row.get("prix_presentation"))
            # Reconstitue prix_m2 quand on a surface + valeur/présentation
            if not prix_m2 and valeur and surface:
                prix_m2 = round(valeur / surface)
            if not prix_m2 and pres and surface:
                prix_m2 = round(pres / 1.07 / surface)
            # Reconstitue surface quand on a prix_m2 + valeur/présentation
            if not surface and prix_m2 and valeur:
                surface = round(valeur / prix_m2)
            if not surface and prix_m2 and pres:
                surface = round(pres / 1.07 / prix_m2)
            # On accepte l'estimation dès qu'on a AU MOINS un signal de prix.
            # Sans surface ni prix_m2, le Classeur affichera un tiret mais le
            # prix de présentation reste visible.
            if not (prix_m2 or valeur or pres):
                skipped += 1
                continue

            # Année : plus récent millésime trouvé
            yrs = _re.findall(r"(19\d\d|20\d\d)", row.get("annee") or "")
            annee = max(yrs) if yrs else ""

            # Quartier : normalise si connu, sinon on garde brut (mais tronqué)
            quartier = (row.get("quartier") or "").strip()
            # Enlève les précisions entre parenthèses et après la première virgule
            quartier = _re.sub(r"\s*\([^)]*\)", "", quartier).split(",")[0].strip()[:120]

            # Description : on préfixe avec le prix LP quand la surface manque
            # (dans ce cas la valeur vénale calculée sera 0)
            description = (row.get("description") or "").strip()
            if not surface or not prix_m2:
                bits = []
                if valeur: bits.append(f"Valeur vénale retenue par LP : CHF {int(valeur):,}".replace(",", "'"))
                if pres: bits.append(f"Prix de présentation : CHF {int(pres):,}".replace(",", "'"))
                if bits:
                    prefix = " · ".join(bits) + "."
                    description = prefix + ("\n" + description if description else "")

            e = Estimation(
                user_id=uid,
                address=adresse,
                quartier=quartier,
                type_bien=type_bien,
                surface=surface or 0,
                annee=annee[:20],
                description=description or None,
                prix_m2=prix_m2 or 0,
                marge=0.07,
                notes=IMPORT_LP_NOTE,
                courtier="Leonard Properties SA",
            )
            db.session.add(e)
            inserted += 1
    db.session.commit()
    return jsonify({"inserted": inserted, "skipped": skipped})


@app.route("/settings", methods=["GET", "POST"])
@login_required
def settings():
    """Paramètres — téléchargement du logo."""
    current_logo = get_logo_path()

    if request.method == "POST":
        # Vérifie s'il y a un fichier dans la requête
        if "logo_file" not in request.files:
            flash("Aucun fichier sélectionné.", "error")
            return redirect(url_for("settings"))

        file = request.files["logo_file"]

        if file.filename == "":
            flash("Aucun fichier sélectionné.", "error")
            return redirect(url_for("settings"))

        # Vérifie l'extension et la taille
        if not allowed_file(file.filename):
            flash("Format non autorisé. Accepté: PNG, JPG, GIF.", "error")
            return redirect(url_for("settings"))

        if file.content_length and file.content_length > 2 * 1024 * 1024:
            flash("Fichier trop volumineux. Max 2 MB.", "error")
            return redirect(url_for("settings"))

        try:
            # Génère un nom de fichier sécurisé avec timestamp
            ext = file.filename.rsplit(".", 1)[1].lower()
            filename = f"logo_{int(datetime.utcnow().timestamp())}.{ext}"
            filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)

            # Supprime l'ancien logo si présent
            if current_logo:
                old_path = os.path.join(app.config["UPLOAD_FOLDER"], current_logo)
                try:
                    os.remove(old_path)
                except OSError:
                    pass

            # Sauvegarde le nouveau fichier
            file.save(filepath)

            # Metà jour la base de données
            setting = Setting.query.filter_by(key="logo_filename").first()
            if setting:
                setting.value = filename
            else:
                setting = Setting(key="logo_filename", value=filename)
                db.session.add(setting)
            db.session.commit()

            flash("Logo mis à jour avec succès.", "success")
            return redirect(url_for("settings"))

        except Exception as e:
            flash(f"Erreur lors de l'upload: {str(e)}", "error")
            return redirect(url_for("settings"))

    return render_template("settings.html", current_logo=current_logo)


# ----------------------- HELPERS -----------------------
def _addr_match(adresse, r):
    if not adresse or not r.adresse:
        return False
    a = adresse.lower().strip()
    b = r.adresse.lower().strip()
    key = a[:14]
    return key in b or b[:14] in a


def ref_for(quartier, adresse, surface=None, type_bien=None):
    """Legacy — retourne un pool d'objets RefPrice (utilisé par les vieux appelants).
    Voir find_comparables() pour la nouvelle version enrichie.
    """
    refs = RefPrice.query.all()
    by_addr = [r for r in refs if _addr_match(adresse, r)]
    by_quartier = [r for r in refs if quartier and r.quartier and
                   r.quartier.lower() == quartier.lower()]
    pool, seen = [], set()
    for r in by_addr + by_quartier:
        if r.id not in seen:
            seen.add(r.id)
            pool.append(r)

    def median(vals):
        vals = sorted(vals)
        return vals[len(vals) // 2] if vals else None

    addr_sold = [r.prix_m2 for r in by_addr if r.prix_m2 and r.kind in ("sold", "retenu")]
    quartier_sold = [r.prix_m2 for r in by_quartier if r.prix_m2 and r.kind in ("sold", "retenu")]
    proposed = median(addr_sold) or median(quartier_sold)
    return proposed, pool


def _same_type_category(a, b):
    """Deux types tombent-ils dans la même famille (appartements OU maisons) ?"""
    if not a or not b:
        return False
    a, b = a.strip(), b.strip()
    if a == b:
        return True
    apt = APARTMENT_TYPES  # défini plus bas dans le fichier
    house = HOUSE_TYPES
    return (a in apt and b in apt) or (a in house and b in house)


NEIGHBOR_QUARTIERS = {
    "Champel": ["Miremont", "Florissant", "Malagnou"],
    "Miremont": ["Champel", "Florissant"],
    "Florissant": ["Champel", "Malagnou", "Miremont"],
    "Malagnou": ["Florissant", "Champel"],
    "Eaux-Vives": ["Villereuse", "La Grange", "Frontenex", "Contamines"],
    "Villereuse": ["Eaux-Vives"],
    "Cologny": ["Vandoeuvres", "Collonge-Bellerive", "Vésenaz"],
    "Vandoeuvres": ["Cologny", "Chêne-Bougeries"],
    "Chêne-Bougeries": ["Chêne-Bourg", "Thônex", "Cologny", "Vandoeuvres"],
    "Chêne-Bourg": ["Chêne-Bougeries", "Thônex"],
    "Thônex": ["Chêne-Bougeries", "Chêne-Bourg"],
    "Vésenaz": ["Collonge-Bellerive", "Cologny"],
    "Collonge-Bellerive": ["Vésenaz", "Anières", "Cologny"],
    "Anières": ["Hermance", "Collonge-Bellerive", "Corsier"],
    "Hermance": ["Anières", "Corsier"],
    "Corsier": ["Anières", "Hermance"],
    "Carouge": ["Plainpalais", "Acacias", "Bâtie"],
    "Plainpalais": ["Jonction", "Acacias", "Cité", "Carouge"],
    "Jonction": ["Plainpalais", "Acacias"],
    "Acacias": ["Carouge", "Plainpalais", "Bâtie"],
    "Pâquis": ["Saint-Gervais", "Grottes", "Nations"],
    "Saint-Gervais": ["Pâquis", "Grottes"],
    "Grottes": ["Pâquis", "Saint-Gervais", "Servette"],
    "Servette": ["Saint-Jean", "Charmilles", "Grottes"],
    "Saint-Jean": ["Servette", "Charmilles"],
}

MIN_COMPARABLES = 3   # objectif minimum


def find_comparables(quartier, address, surface, type_bien, current_eid=None, user_id=None):
    """Trouve les comparables réels pour un bien avec fallback progressif.

    Ordre :
      1. Filtre strict — même quartier, surface ±20 %, même famille de type
      2. Élargir surface à ±40 % si <3 résultats
      3. Ignorer la surface si <3 résultats
      4. Inclure les quartiers voisins si <3 résultats
      5. Ignorer le type si toujours <3 résultats

    Retourne (proposed_prix_m2, comparables_dicts, stats) où stats.match_level indique
    quel niveau de filtrage a été effectivement utilisé.
    """
    def _search(_quartiers, _surface_tol, _use_type):
        return _do_search(_quartiers, surface, _surface_tol, type_bien if _use_type else None,
                          current_eid, user_id)

    # Niveau 1 : strict
    comps = _search([quartier], 0.20, True)
    match_level = "strict"

    # Niveau 2 : élargir surface à ±40 %
    if len(comps) < MIN_COMPARABLES:
        comps = _search([quartier], 0.40, True)
        match_level = "surface ±40 %"

    # Niveau 3 : ignorer surface
    if len(comps) < MIN_COMPARABLES:
        comps = _search([quartier], None, True)
        match_level = "quartier + type"

    # Niveau 4 : ajouter quartiers voisins
    if len(comps) < MIN_COMPARABLES and quartier:
        neighbors = [quartier] + NEIGHBOR_QUARTIERS.get(quartier, [])
        comps = _search(neighbors, None, True)
        match_level = "quartier + voisins"

    # Niveau 5 : tout le quartier + voisins, tous types
    if len(comps) < MIN_COMPARABLES and quartier:
        neighbors = [quartier] + NEIGHBOR_QUARTIERS.get(quartier, [])
        comps = _search(neighbors, None, False)
        match_level = "voisins, tous types"

    # Tri : comparables avec description riche d'abord (utile pour le rapport),
    # puis année la plus récente, puis LP > vendus > retenus > à vendre.
    kind_order = {"estimation": 0, "sold": 1, "retenu": 2, "forsale": 3}
    def _year_of(c):
        try:
            return int(str(c.get("annee") or "0")[:4])
        except ValueError:
            return 0
    comps.sort(key=lambda c: (
        0 if c.get("description") else 1,           # descriptions d'abord
        -_year_of(c),                                # année desc
        kind_order.get(c["kind"], 9),                # LP → sold → retenu → forsale
    ))

    # Dédoublonnage : même adresse + surface + prix/m² → on garde la 1re occurrence
    # (celle avec description ou la plus récente d'après le tri qui précède).
    def _addr_key(a):
        return " ".join((a or "").lower().split()).rstrip(",.")
    seen = set()
    unique = []
    for c in comps:
        key = (_addr_key(c.get("adresse")),
               int(c["surface"] or 0),
               int(c["prix_m2"] or 0))
        if key in seen:
            continue
        seen.add(key)
        unique.append(c)
    comps = unique

    # Prix proposé = moyenne des prix/m² des comparables vendus / LP / retenus
    sold_pm2 = [c["prix_m2"] for c in comps if c["kind"] in ("sold", "retenu", "estimation")]
    forsale_pm2 = [c["prix_m2"] for c in comps if c["kind"] == "forsale"]
    proposed = None
    if sold_pm2:
        proposed = round(sum(sold_pm2) / len(sold_pm2))
    elif forsale_pm2:
        proposed = round(sum(forsale_pm2) / len(forsale_pm2) * 0.95)

    stats = {
        "total": len(comps),
        "n_sold": sum(1 for c in comps if c["kind"] in ("sold", "retenu", "estimation")),
        "n_forsale": sum(1 for c in comps if c["kind"] == "forsale"),
        "min_pm2": min((c["prix_m2"] for c in comps), default=None),
        "max_pm2": max((c["prix_m2"] for c in comps), default=None),
        "avg_pm2": proposed,
        "match_level": match_level,
    }
    return proposed, comps, stats


def _do_search(quartiers, surface, surface_tol, type_bien, current_eid=None, user_id=None):
    """Recherche interne — ne fait que le filtre correspondant aux paramètres passés."""
    comparables = []
    quartiers_lc = {q.lower() for q in quartiers if q}

    def _quartier_ok(r_quartier):
        if not quartiers_lc:
            return True
        return r_quartier and r_quartier.lower() in quartiers_lc

    def _surface_ok(r_surface):
        if surface_tol is None or surface is None:
            return True
        if not r_surface:
            return False
        lo, hi = surface * (1 - surface_tol), surface * (1 + surface_tol)
        return lo <= r_surface <= hi

    def _type_ok(r_type):
        if not type_bien:
            return True
        if not r_type:
            return False
        return _same_type_category(r_type, type_bien)

    # ---- Source 1 : estimations passées de l'utilisateur ----
    if user_id:
        q = Estimation.query.filter_by(user_id=user_id)
        if current_eid:
            q = q.filter(Estimation.id != current_eid)
        for est in q.all():
            if not _quartier_ok(est.quartier):
                continue
            if not _type_ok(est.type_bien):
                continue
            if not _surface_ok(est.surface):
                continue
            if not est.prix_m2 or not est.surface:
                continue
            comparables.append({
                "adresse": est.address or "—",
                "quartier": est.quartier or "",
                "type_bien": est.type_bien or "",
                "surface": est.surface,
                "prix_m2": est.prix_m2,
                "prix_total": est.prix_presentation,
                "annee": est.annee or "",
                "etat": est.etat or "",
                "atouts": est.atouts or "",
                "inconvenients": est.inconvenients or "",
                "description": est.description or "",
                "reference": "",
                "kind": "estimation",
                "source": "Estimation LP",
                "date": est.created_at.strftime("%m/%Y") if est.created_at else "",
            })

    # ---- Source 2 : RefPrice (marché + CSV LP) ----
    for r in RefPrice.query.all():
        if not _quartier_ok(r.quartier):
            continue
        if not _type_ok(r.type_bien):
            continue
        if not _surface_ok(r.surface):
            continue
        if not r.prix_m2:
            continue
        prix_total = (r.prix_m2 * r.surface) if r.surface else None
        comparables.append({
            "adresse": r.adresse or "—",
            "quartier": r.quartier or "",
            "type_bien": r.type_bien or "",
            "surface": r.surface,
            "prix_m2": r.prix_m2,
            "prix_total": r.prix_total or prix_total,
            "annee": r.annee or "",
            "etat": "",
            "atouts": "",
            "inconvenients": "",
            "description": r.description or "",
            "reference": r.reference or "",
            "kind": r.kind or "sold",
            "source": r.source or "Marché",
            "date": r.annee or "",
        })

    return comparables


# ----------------------- ROUTES -----------------------
@app.route("/")
def index():
    if session.get("user_id"):
        return redirect(url_for("dashboard"))
    return render_template("landing.html")


@app.route("/upload", methods=["POST"])
@login_required
def upload():
    """Reçoit un fichier déposé, en extrait les champs, pré-remplit le formulaire."""
    file = request.files.get("file")
    if not file or not file.filename:
        return jsonify({"error": "Aucun fichier"}), 400
    text = extract_text(file)
    fields = extract_fields(text)
    proposed, pool = ref_for(fields.get("quartier"), fields.get("address"))
    if proposed and not fields.get("prix_m2"):
        fields["prix_m2"] = round(proposed)
    fields["_refs"] = [{"adresse": r.adresse, "quartier": r.quartier,
                        "prix_m2": r.prix_m2, "annee": r.annee, "kind": r.kind} for r in pool]
    return jsonify(fields)


class PhotoExtraction(BaseModel):
    """Champs extraits d'une photo, d'un texte libre ou d'une fiche/annonce d'un bien à Genève."""
    address: str = Field(default="", description="Adresse complète (ex: Avenue de Miremont 30)")
    quartier: str = Field(default="", description="Quartier ou commune genevoise (ex: Champel, Eaux-Vives, Miremont, Cologny, Chêne-Bougeries, Vésenaz, Carouge, Pâquis, Servette…). Vide si vraiment inconnu.")
    type_bien: str = Field(default="", description="Type — Appartement, Duplex, Attique, Penthouse, Triplex, Maison individuelle, Villa")
    surface: Optional[float] = Field(default=None, description="Surface habitable ou pondérée en m² (nombre uniquement, sans unité)")
    pieces: str = Field(default="", description="Nombre de pièces (ex: 5, 4.5)")
    etage: str = Field(default="", description="Étage (ex: 3, RDC, dernier)")
    annee: str = Field(default="", description="Année de construction ou de rénovation (ex: 2016)")
    etat: str = Field(default="", description="État du bien (ex: Excellent, Rénové 2024, À rafraîchir)")
    balcon: Optional[float] = Field(default=None, description="Surface totale des extérieurs (balcon/loggia/terrasse/jardin) en m²")
    parking_nb: Optional[int] = Field(default=None, description="Nombre de places de parking")
    description: str = Field(default="", description="Description libre du bien (2-4 phrases max, style LP)")


def _extraction_prompt() -> str:
    quartiers_hint = ", ".join(q for q in QUARTIERS if q != "Autre")
    return (
        "Extrais les champs demandés pour pré-remplir un formulaire d'estimation Leonard Properties. "
        "Règles :\n"
        f"- 'quartier' : choisis dans cette liste si possible : {quartiers_hint}. "
        "Si le bien est ailleurs à Genève, mets le nom de la commune/quartier tel qu'il apparaît. "
        "Vide seulement si vraiment inconnu.\n"
        "- 'surface' : nombre en m² uniquement (sans unité). Si 'surface pondérée' et 'surface PPE' "
        "apparaissent, prends la pondérée.\n"
        "- 'description' : 3 à 5 phrases en français, style sobre et factuel LP. Mentionne le quartier "
        "et un atout crédible. N'invente rien.\n"
        "- Si un champ n'est pas mentionné, laisse-le vide. Ne devine pas — mieux vaut vide que faux."
    )


def _extract_from_content(content_blocks) -> dict:
    """Appelle Claude Opus 4.8 avec un content (image ou texte) → dict de champs prêt pour le formulaire."""
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY manquante")
    from anthropic import Anthropic
    client = Anthropic(api_key=api_key)
    response = client.messages.parse(
        model="claude-opus-4-8",
        max_tokens=1500,
        messages=[{"role": "user", "content": content_blocks}],
        output_format=PhotoExtraction,
    )
    fields = response.parsed_output.model_dump()
    return {k: ("" if v is None else v) for k, v in fields.items()}


def _resize_image_for_vision(raw: bytes, max_edge: int = 1568) -> tuple[bytes, str]:
    """Redimensionne + reencode l'image pour Claude Vision. Retourne (bytes, media_type)."""
    from PIL import Image, ImageOps
    img = Image.open(io.BytesIO(raw))
    img = ImageOps.exif_transpose(img)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    w, h = img.size
    if max(w, h) > max_edge:
        img.thumbnail((max_edge, max_edge))
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    return buf.getvalue(), "image/jpeg"


@app.route("/upload-photo", methods=["POST"])
@login_required
def upload_photo():
    """Analyse une photo (fiche, capture d'écran, photo du bien) via Claude Vision et pré-remplit le formulaire."""
    file = request.files.get("file")
    if not file or not file.filename:
        return jsonify({"error": "Aucun fichier"}), 400

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        return jsonify({"error": "Clé API Anthropic non configurée sur le serveur (ANTHROPIC_API_KEY manquante)."}), 500

    try:
        raw = file.read()
        img_bytes, media_type = _resize_image_for_vision(raw)
    except Exception as e:
        return jsonify({"error": f"Image illisible (format non supporté ou fichier corrompu) : {e}"}), 400

    b64 = base64.standard_b64encode(img_bytes).decode()

    try:
        fields = _extract_from_content([
            {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64}},
            {"type": "text", "text": "Cette image montre une fiche, une annonce, une capture d'écran ou une photo d'un bien immobilier à Genève. " + _extraction_prompt()},
        ])
    except Exception as e:
        return jsonify({"error": f"Erreur lors de l'analyse par l'IA : {e}"}), 502

    proposed, pool = ref_for(fields.get("quartier"), fields.get("address"))
    if proposed:
        fields["prix_m2"] = round(proposed)
    fields["_refs"] = [{"adresse": r.adresse, "quartier": r.quartier,
                        "prix_m2": r.prix_m2, "annee": r.annee, "kind": r.kind} for r in pool]
    return jsonify(fields)


@app.route("/analyze-text", methods=["POST"])
@login_required
def analyze_text():
    """Analyse une description libre du bien collée par l'utilisateur → pré-remplit tous les champs."""
    data = request.get_json(silent=True) or request.form
    text = (data.get("text") or "").strip()
    if not text:
        return jsonify({"error": "Texte vide."}), 400
    if len(text) < 15:
        return jsonify({"error": "Décris le bien avec un peu plus de détails (au moins une phrase)."}), 400

    if not os.environ.get("ANTHROPIC_API_KEY"):
        return jsonify({"error": "Clé API Anthropic non configurée sur le serveur."}), 500

    try:
        fields = _extract_from_content([
            {"type": "text", "text": (
                "Voici une description libre d'un bien immobilier à Genève, écrite par un courtier "
                "Leonard Properties ou collée depuis une annonce. " + _extraction_prompt() +
                "\n\n--- Description ---\n" + text
            )},
        ])
    except Exception as e:
        return jsonify({"error": f"Erreur lors de l'analyse par l'IA : {e}"}), 502

    proposed, pool = ref_for(fields.get("quartier"), fields.get("address"))
    if proposed:
        fields["prix_m2"] = round(proposed)
    fields["_refs"] = [{"adresse": r.adresse, "quartier": r.quartier,
                        "prix_m2": r.prix_m2, "annee": r.annee, "kind": r.kind} for r in pool]
    return jsonify(fields)


QUARTIER_ATOUTS = {
    "Champel": "quartier résidentiel prisé de la rive gauche, écoles internationales et hôpitaux à proximité, ambiance calme et verte",
    "Eaux-Vives": "quartier vivant entre le lac et le parc La Grange, commerces et restaurants nombreux, très bien desservi",
    "Miremont": "secteur résidentiel de Champel-Miremont, calme, familial, vues sur la ville et proximité des parcs",
}


def _generate_lp_description(fields: dict) -> str:
    """Génère un paragraphe LP (3-5 phrases) via Claude Opus 4.8. Retourne '' si l'API n'est pas dispo."""
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        return ""
    try:
        from anthropic import Anthropic
        client = Anthropic(api_key=api_key)
        atouts_quartier = QUARTIER_ATOUTS.get(fields.get("quartier") or "", "")
        prompt = f"""Rédige un paragraphe de 3 à 5 phrases décrivant ce bien immobilier dans le style sobre, factuel et élégant de Leonard Properties (courtier à Genève).

Consignes :
- Français impeccable, phrases complètes, ton professionnel LP.
- Évite les superlatifs vides ("magnifique", "exceptionnel", "unique") sauf s'ils sont justifiés par les faits.
- Décris le bien, mentionne le quartier et ses atouts, cite l'état/année si pertinent.
- Ne mentionne pas de champ vide ou inconnu.
- Pas de listes à puces, pas d'en-têtes, uniquement un ou deux paragraphes courts.
- N'invente rien qui ne soit pas dans les données ci-dessous.

Bien à décrire :
- Adresse : {fields.get('address') or 'non renseignée'}
- Quartier : {fields.get('quartier') or 'non renseigné'}{f' ({atouts_quartier})' if atouts_quartier else ''}
- Type : {fields.get('type_bien') or 'non renseigné'}
- Surface : {fields.get('surface') or '?'} m²
- Pièces : {fields.get('pieces') or '?'}
- Étage : {fields.get('etage') or '?'}
- Année : {fields.get('annee') or '?'}
- État : {fields.get('etat') or '?'}
- Extérieurs : {fields.get('balcon') or 0} m²
- Parkings : {fields.get('parking_nb') or 0}

Réponds uniquement avec le paragraphe (pas de préambule, pas de guillemets)."""
        response = client.messages.create(
            model="claude-opus-4-8",
            max_tokens=500,
            messages=[{"role": "user", "content": prompt}],
        )
        text = "".join(b.text for b in response.content if b.type == "text").strip()
        return text
    except Exception as e:
        app.logger.warning(f"Description IA échouée : {e}")
        return ""


def _format_comparables_for_prompt(comparables):
    """Formate 5 comparables pour le prompt Claude — reste concis."""
    lines = []
    for c in comparables[:5]:
        parts = [c.get("adresse") or "—"]
        if c.get("type_bien"):
            parts.append(c["type_bien"])
        if c.get("surface"):
            parts.append(f"{c['surface']} m²")
        if c.get("annee"):
            parts.append(c["annee"])
        if c.get("prix_m2"):
            parts.append(f"{int(c['prix_m2'])} CHF/m²")
        if c.get("prix_total"):
            parts.append(f"prix total ≈ {int(c['prix_total'])} CHF")
        line = " · ".join(str(p) for p in parts)
        if c.get("description"):
            desc = " ".join(c["description"].split())[:200]
            line += f" — {desc}"
        lines.append(f"- {line}")
    return "\n".join(lines) if lines else "(aucun comparable retenu)"


def _generate_lp_report(e, comparables, stats, force=False):
    """Génère un rapport LP complet via Claude Opus 4.8 et le renvoie sous forme HTML.
    Résultat mis en cache dans e.report_ai. Si force=True, force la régénération.
    """
    if not force and e.report_ai:
        return e.report_ai

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        return ""

    try:
        from anthropic import Anthropic
        client = Anthropic(api_key=api_key)

        atouts_quartier = QUARTIER_ATOUTS.get(e.quartier or "", "")
        comp_block = _format_comparables_for_prompt(comparables or [])
        match_level = (stats or {}).get("match_level", "strict")
        avg_pm2 = (stats or {}).get("avg_pm2")
        min_pm2 = (stats or {}).get("min_pm2")
        max_pm2 = (stats or {}).get("max_pm2")

        is_house = e.type_bien in HOUSE_TYPES

        prompt = f"""Tu rédiges le corps d'un rapport d'estimation Leonard Properties SA (courtier immobilier à Genève).
Tu dois **respecter STRICTEMENT la méthode, la structure et la syntaxe officielle des rapports LP** décrites ci-dessous.

DONNÉES DU BIEN :
- Adresse : {e.address or 'non précisée'}
- Quartier / commune : {e.quartier or '—'}{f' — atouts : {atouts_quartier}' if atouts_quartier else ''}
- Type : {e.type_bien or '—'}
- Surface habitable : {e.surface or '?'} m²
- Pièces : {e.pieces or '?'} · Étage : {e.etage or '?'} · Année : {e.annee or '?'}
- État : {e.etat or '?'}
- Extérieurs : {e.balcon or 0} m² (pondération {int((e.balcon_pond or 0) * 100)} %)
- Parkings : {e.parking_nb or 0} (valeur unitaire {int(e.parking_val or 0)} CHF)
- Prix/m² retenu : {int(e.prix_m2 or 0)} CHF
- Description saisie par le courtier : {(e.description or '').strip() or '—'}
- Atouts renseignés : {(e.atouts or '').strip() or '—'}
- Points d'attention renseignés : {(e.inconvenients or '').strip() or '—'}
- Marge négociation + commission : {int((e.marge or 0) * 100)} %
- Valeur vénale calculée : {int(e.valeur_venale or 0)} CHF
- Prix de présentation calculé : {int(e.prix_presentation or 0)} CHF

COMPARABLES DU PANEL (filtre appliqué : {match_level}) :
{comp_block}
Statistiques du panel : moyenne {avg_pm2 or '—'} CHF/m² · fourchette {min_pm2 or '—'} – {max_pm2 or '—'} CHF/m².

============================================================
MÉTHODE OFFICIELLE LEONARD PROPERTIES — À RESPECTER
============================================================

**Définition LP de la valeur vénale (à reprendre presque verbatim) :**
« La valeur vénale d'un bien immobilier représente le prix que l'on peut raisonnablement espérer obtenir lors de sa vente dans un délai d'un an, dans des conditions normales de marché. »

**Méthode hédoniste LP (à reprendre presque verbatim) :**
« Nous appliquons la méthode hédoniste pour l'évaluation de {'maisons individuelles' if is_house else 'appartements en PPE'}. Cette méthode comparative s'appuie sur un historique des ventes que nous tenons à jour de biens comparables. Les paramètres déterminants incluent : l'emplacement, la qualité de construction, l'architecture, le coefficient d'impôt, la superficie {'du terrain, ' if is_house else ''}le vis-à-vis, le voisinage direct, l'exposition et la vue. Nos experts évaluent ces différents paramètres en fonction de leur importance pour déterminer la valeur vénale. »

{'**Méthode complémentaire pour la villa — valeur intrinsèque :** cite brièvement que LP calcule aussi la valeur réelle (coût terrain × surface parcelle + coût de construction) à titre indicatif, et que l écart entre intrinsèque et hédoniste renseigne sur le dynamisme du marché local.' if is_house else ''}

**Formulations LP à reprendre :**
- Pour introduire le prix retenu : « En tenant compte des caractéristiques des propriétés similaires vendues et à vendre, et de la moyenne des prix au m² obtenue, nous retenons un prix au m² de CHF {int(e.prix_m2 or 0):,} / m². »
- Pour conclure : « Ce prix serait considéré comme un prix de présentation initial ambitieux et réaliste, qui nous permettrait de jauger le marché pour ensuite procéder à une adaptation si nécessaire. »
- Réserves LP standards : « Tout élément ou fait qui n'aurait pas été porté à la connaissance de l'expert et qui pourrait modifier son appréciation, est dûment réservé. Cette estimation n'est pas une expertise foncière officielle et un écart de +/- 10 % est possible selon les conditions de marché. »

**Amortissement :** si le bien n'a pas été rénové depuis plus de 10 ans, mentionner qu'un amortissement de 1 % par an peut être appliqué pour tenir compte de la vétusté (méthode LP).

============================================================
STRUCTURE OFFICIELLE À GÉNÉRER (dans cet ordre EXACT)
============================================================

1. **`<h2>1. Introduction</h2>`** — Une lettre courtoise ouvrant sur « Cher Monsieur, Chère Madame, », rappelant le contexte de la mission (« Dans le prolongement de notre visite… ») et annonçant que ce rapport présente l'estimation de valeur vénale ainsi qu'une stratégie de commercialisation. 3–4 phrases.

2. **`<h2>2. Méthode d'évaluation</h2>`** — Reprends la définition LP de la valeur vénale (verbatim ci-dessus), puis explicite la méthode hédoniste comparative telle que définie ci-dessus. {'Ajoute un paragraphe sur la valeur intrinsèque en complément pour cette maison / villa.' if is_house else ''}

3. **`<h2>3. Situation, zone et lois</h2>`** — Situe le bien à l'échelle macro (Genève, quartier) puis micro (accès, transports, écoles, commerces). Si possible, cite la zone d'affectation vraisemblable (Zone 1 vieille-ville, Zone 3 urbaine dense, Zone 5 résidentielle villas, etc.) et mentionne les implications légales génériques en une phrase. Reste factuel.

4. **`<h2>4. Descriptif du bien</h2>`** — Sous-sections `<h3>Situation macro et micro</h3>`, `<h3>Descriptif détaillé</h3>` (distribution, matériaux, exposition, vues), `<h3>Prestations et état général</h3>`, `<h3>Travaux à prévoir</h3>` si pertinent. Riche mais factuel.

5. **`<h2>5. Atouts et points d'attention</h2>`** — Deux sous-sections `<h3>Atouts</h3><ul>…</ul>` et `<h3>Points d'attention</h3><ul>…</ul>` avec des puces courtes, précises, non redondantes.

6. **`<h2>6. Estimation hédoniste — analyse des comparables</h2>`** — Deux ou trois paragraphes qui commentent le panel : qualité et récence des transactions retenues, fourchette de prix/m² observée, positionnement du bien vis-à-vis de ce panel (au-dessus / dans / en-dessous et pourquoi). **Cite au moins 2 comparables par leur adresse.** Termine sur la phrase LP standard « En tenant compte des caractéristiques des propriétés similaires vendues et à vendre, et de la moyenne des prix au m² obtenue, nous retenons un prix au m² de CHF {int(e.prix_m2 or 0):,} / m². »

7. **`<h2>7. Estimation de la valeur vénale</h2>`** — Rappelle la formule : `prix/m² × surface + prix/m² × extérieurs pondérés + parkings`. Détaille les 2–3 composantes en une petite liste ou un paragraphe. Si l'année de construction est ancienne, mentionne l'éventuel amortissement pour vétusté. Termine sur : « **TOTAL VALEUR VÉNALE : CHF {int(e.valeur_venale or 0):,}.-** ».

8. **`<h2>8. Valeur vénale et prix de présentation</h2>`** — Petit paragraphe qui affiche :
   `<p>Valeur de marché (méthode hédoniste comparative) : <strong>CHF {int(e.valeur_venale or 0):,}.-</strong></p>`
   `<p>Prix de présentation suggéré (marge de négociation et commission incluses) : <strong>CHF {int(e.prix_presentation or 0):,}.-</strong></p>`
   Puis reprends la conclusion LP verbatim (« Ce prix serait considéré comme un prix de présentation initial ambitieux et réaliste… »).

9. **`<h2>9. Conclusions et réserves</h2>`** — Reprends les réserves LP verbatim (paragraphe ci-dessus), puis une phrase de clôture courtoise (« Nous restons à votre entière disposition pour tout complément d'information et serions ravis de collaborer avec vous. »).

============================================================
CONSIGNES DE FORME
============================================================
- **Ton LP** : professionnel, sobre, factuel, français impeccable, phrases complètes. Le vouvoiement est de rigueur.
- **Pas de superlatifs vides** (« magnifique », « exceptionnel », « unique ») sauf si les faits les justifient.
- **N'invente rien** au-delà des données fournies ; si un champ est vide, ne le mentionne pas.
- **HTML strict** : n'utilise que `<h2>`, `<h3>`, `<p>`, `<ul>`, `<li>`, `<strong>`, `<em>`. Pas de `<div>`, pas de style inline, pas de titre H1.
- **Format** : réponds uniquement avec le HTML des 9 sections dans l'ordre, sans préambule, sans markdown, sans triple-backtick."""

        response = client.messages.create(
            model="claude-sonnet-5",
            max_tokens=2500,
            messages=[{"role": "user", "content": prompt}],
        )
        html = "".join(b.text for b in response.content if b.type == "text").strip()
        # Nettoyage : retire d'éventuels blocs Markdown
        if html.startswith("```"):
            html = html.split("```", 2)[1]
            if html.startswith("html"):
                html = html[4:]
            html = html.strip()
        # Persiste
        e.report_ai = html
        db.session.commit()
        return html
    except Exception as exc:
        app.logger.warning(f"Rapport IA échoué : {exc}")
        return ""


@app.route("/admin/purge-duplicates")
@login_required
def admin_purge_duplicates():
    """Supprime les estimations en doublon (même adresse+surface+prix/m²).
    On garde la plus récente pour chaque clé."""
    uid = session.get("user_id")
    seen = {}
    to_delete = []
    # ⚠️ On ne dédoublonne QUE les imports LP entre eux — jamais les créations
    # manuelles de l'utilisateur, sinon on risque de supprimer une estimation
    # qu'il vient de créer parce qu'un import LP a la même adresse.
    ests = (Estimation.query
            .filter_by(user_id=uid, notes=IMPORT_LP_NOTE)
            .order_by(Estimation.created_at.desc())
            .all())
    for e in ests:
        key = " ".join((e.address or "").lower().split()).rstrip(",.")
        if not key:
            continue
        if key in seen:
            to_delete.append(e)
        else:
            seen[key] = e
    for e in to_delete:
        db.session.delete(e)
    db.session.commit()
    return jsonify({"deleted": len(to_delete), "kept_lp_imports": len(seen)})


@app.route("/api/preview-comparables", methods=["POST"])
@login_required
def api_preview_comparables():
    """Renvoie un aperçu léger des comparables pour le formulaire de saisie."""
    data = request.get_json(silent=True) or {}
    quartier = (data.get("quartier") or "").strip()
    type_bien = (data.get("type_bien") or "").strip()
    address = (data.get("address") or "").strip()
    try:
        surface = float(str(data.get("surface") or "").replace("'", "").replace(" ", "") or 0)
    except ValueError:
        surface = 0
    if not quartier and not address:
        return jsonify({"comparables": [], "stats": None, "proposed_pm2": None})
    proposed_pm2, comps, stats = find_comparables(
        quartier, address, surface, type_bien,
        current_eid=None, user_id=session.get("user_id"),
    )
    # Sur la preview, on écarte les estimations avec la même adresse que celle en
    # cours de saisie (souvent des essais précédents dupliqués)
    if address:
        key = " ".join(address.lower().split()).rstrip(",.")
        def _same_addr(a):
            return " ".join((a or "").lower().split()).rstrip(",.") == key
        comps = [c for c in comps if not _same_addr(c.get("adresse"))]
    # Version allégée pour le formulaire (max 6)
    lite = [{
        "adresse": c["adresse"],
        "type_bien": c.get("type_bien"),
        "annee": c.get("annee"),
        "surface": c.get("surface"),
        "prix_total": c.get("prix_total"),
        "prix_m2": c.get("prix_m2"),
        "kind": c.get("kind"),
        "source": c.get("source"),
    } for c in comps[:6]]
    return jsonify({
        "comparables": lite,
        "total": len(comps),
        "proposed_pm2": proposed_pm2,
        "match_level": stats.get("match_level") if stats else None,
    })


@app.route("/estimation/new", methods=["GET", "POST"])
@login_required
def estimation_new():
    if request.method == "POST":
        f = request.form
        def num(k, d=0.0):
            try:
                return float(str(f.get(k, "")).replace("'", "").replace(" ", "").replace(",", ".") or d)
            except ValueError:
                return d
        description = (f.get("description") or "").strip()
        # Si l'utilisateur n'a rien mis, on demande à Claude de rédiger un paragraphe LP
        if not description:
            description = _generate_lp_description({
                "address": f.get("address"), "quartier": f.get("quartier"),
                "type_bien": f.get("type_bien"), "surface": num("surface"),
                "pieces": f.get("pieces"), "etage": f.get("etage"),
                "annee": f.get("annee"), "etat": f.get("etat"),
                "balcon": num("balcon"), "parking_nb": int(num("parking_nb")),
            })
        e = Estimation(
            user_id=session.get("user_id"),
            address=f.get("address"), quartier=f.get("quartier"),
            type_bien=f.get("type_bien"), surface=num("surface"),
            pieces=f.get("pieces"), etage=f.get("etage"), annee=f.get("annee"),
            etat=f.get("etat"), balcon=num("balcon"), balcon_pond=num("balcon_pond", 0.5),
            parking_nb=int(num("parking_nb")), parking_val=num("parking_val"),
            prix_m2=num("prix_m2"), marge=num("marge", 0.07),
            description=description, atouts=f.get("atouts"),
            inconvenients=f.get("inconvenients"), courtier=f.get("courtier"))
        db.session.add(e)
        db.session.commit()
        return redirect(url_for("estimation_report", eid=e.id))
    return render_template("estimation_form.html", quartiers=QUARTIERS, prefill={})


import threading

# Suivi des générations en cours (id → True) pour ne pas relancer deux fois en parallèle
_generations_in_progress = set()
_generations_lock = threading.Lock()


def _generate_report_in_background(app_ref, eid, force):
    """Thread : régénère le rapport pour l'estimation eid et le persiste."""
    with app_ref.app_context():
        try:
            e = Estimation.query.get(eid)
            if not e:
                return
            _, comparables, stats = find_comparables(
                e.quartier, e.address, e.surface, e.type_bien,
                current_eid=e.id, user_id=e.user_id,
            )
            html = _generate_lp_report(e, comparables, stats, force=force) or ""
            if not html:
                _, legacy_pool = ref_for(e.quartier, e.address)
                sold = [r for r in legacy_pool if r.kind in ("sold", "retenu")]
                forsale = [r for r in legacy_pool if r.kind == "forsale"]
                html = build_report(e, sold, forsale)
            if html:
                e.report_ai = html
                db.session.commit()
        except Exception as exc:
            app_ref.logger.exception(f"Génération background KO pour {eid} : {exc}")
        finally:
            with _generations_lock:
                _generations_in_progress.discard(eid)


@app.route("/estimation/<int:eid>")
@login_required
def estimation_report(eid):
    e = own_estimation_or_404(eid)
    force = request.args.get("regenerate") == "1"

    # Si le rapport est déjà en cache et qu'on ne force pas, on affiche direct
    if e.report_ai and not force:
        try:
            proposed_pm2, comparables, stats = find_comparables(
                e.quartier, e.address, e.surface, e.type_bien,
                current_eid=e.id, user_id=session.get("user_id"),
            )
        except Exception as exc:
            app.logger.exception(f"find_comparables KO pour {eid} : {exc}")
            proposed_pm2, comparables, stats = None, [], {
                "match_level": "erreur", "total": 0, "n_sold": 0, "n_forsale": 0,
                "min_pm2": None, "max_pm2": None, "avg_pm2": None}
        return render_template("report.html", e=e, report_html=e.report_ai,
                               comparables=comparables, stats=stats,
                               proposed_pm2=proposed_pm2, ai_generated=True)

    # Sinon on lance la génération en arrière-plan si pas déjà en cours
    if force:
        e.report_ai = None
        db.session.commit()

    should_start = False
    with _generations_lock:
        if eid not in _generations_in_progress:
            _generations_in_progress.add(eid)
            should_start = True
    if should_start:
        threading.Thread(
            target=_generate_report_in_background,
            args=(app._get_current_object(), eid, force),
            daemon=True,
        ).start()

    # On renvoie une page d'attente qui recharge toute seule
    return render_template("report_generating.html", eid=eid, address=e.address)


@app.route("/estimation/<int:eid>/status")
@login_required
def estimation_status(eid):
    """Endpoint JSON : le rapport est-il prêt ?"""
    e = own_estimation_or_404(eid)
    ready = bool(e.report_ai)
    return jsonify({"ready": ready})


@app.route("/estimation/<int:eid>/delete", methods=["POST"])
@login_required
def estimation_delete(eid):
    e = own_estimation_or_404(eid)
    db.session.delete(e)
    db.session.commit()
    return redirect(url_for("index"))


@app.route("/references")
@login_required
def references():
    refs = RefPrice.query.order_by(RefPrice.quartier, RefPrice.adresse).all()
    return render_template("references.html", refs=refs, quartiers=QUARTIERS)


@app.route("/references/add", methods=["POST"])
@login_required
def references_add():
    f = request.form
    try:
        pm = float(str(f.get("prix_m2", "")).replace("'", "").replace(" ", "") or 0)
    except ValueError:
        pm = 0
    db.session.add(RefPrice(quartier=f.get("quartier"), adresse=f.get("adresse"),
                            prix_m2=pm, annee=f.get("annee"), source=f.get("source"),
                            kind=f.get("kind", "sold")))
    db.session.commit()
    return redirect(url_for("references"))


@app.route("/references/<int:rid>/delete", methods=["POST"])
@login_required
def references_delete(rid):
    r = RefPrice.query.get_or_404(rid)
    db.session.delete(r)
    db.session.commit()
    return redirect(url_for("references"))


@app.route("/analysis")
@login_required
def analysis():
    """Page d'analyse rapide avec texte libre."""
    return render_template("analysis.html")


@app.route("/quick-analyze", methods=["POST"])
@login_required
def quick_analyze():
    """Analyse un texte libre et retourne une estimation complète."""
    data = request.get_json() or {}
    text = data.get("text", "").strip()

    if not text:
        return jsonify({"error": "Texte vide"}), 400

    # Extraire les données du texte
    fields = extract_fields(text)

    # Récupérer les références prix
    proposed, pool = ref_for(fields.get("quartier"), fields.get("address"))
    if proposed and not fields.get("prix_m2"):
        fields["prix_m2"] = round(proposed)

    # Construire l'objet Estimation
    def num(v, d=0.0):
        try:
            if isinstance(v, (int, float)):
                return float(v)
            return float(str(v or d).replace("'", "").replace(" ", "").replace(",", ".") or d)
        except (ValueError, TypeError):
            return d

    est = Estimation(
        address=fields.get("address", ""),
        quartier=fields.get("quartier", ""),
        type_bien=fields.get("type_bien", ""),
        surface=num(fields.get("surface")),
        pieces=fields.get("pieces", ""),
        etage=fields.get("etage", ""),
        annee=fields.get("annee", ""),
        etat=fields.get("etat", ""),
        balcon=num(fields.get("balcon")),
        balcon_pond=num(fields.get("balcon_pond", 0.5), 0.5),
        parking_nb=int(num(fields.get("parking_nb"))),
        parking_val=num(fields.get("parking_val")),
        prix_m2=num(fields.get("prix_m2")),
        marge=num(fields.get("marge", 0.07), 0.07),
        description=fields.get("description", "")
    )

    # Retourner le résultat
    sold = [r for r in pool if r.kind in ("sold", "retenu")]
    return jsonify({
        "address": est.address,
        "quartier": est.quartier,
        "surface": est.surface,
        "condition": est.etat,
        "estimated_price": int(est.prix_presentation),
        "price_per_m2": int(est.prix_m2) if est.prix_m2 else 0,
        "comparables": [
            {
                "address": r.adresse,
                "surface": est.surface,
                "price": int(est.prix_presentation),
                "price_per_m2": int(r.prix_m2)
            } for r in sold[:3]
        ]
    })


# ----------------------- ESTIMATIONS HISTORIQUE & GESTION -----------------------
@app.route("/estimations")
@login_required
def estimations_list():
    """Liste toutes les estimations avec recherche et filtres."""
    q = request.args.get("q", "").strip()
    quartier = request.args.get("quartier", "").strip()

    query = own_estimations()
    if q:
        query = query.filter(
            db.or_(
                Estimation.address.ilike(f"%{q}%"),
                Estimation.quartier.ilike(f"%{q}%")
            )
        )
    if quartier:
        query = query.filter_by(quartier=quartier)

    estimations = query.order_by(Estimation.created_at.desc()).all()
    return render_template("estimations_list.html", estimations=estimations,
                         quartiers=QUARTIERS, current_q=q, current_quartier=quartier)


@app.route("/estimation/<int:eid>/clone", methods=["POST"])
@login_required
def estimation_clone(eid):
    """Clone une estimation existante."""
    e = own_estimation_or_404(eid)
    new_e = Estimation(
        address=e.address, quartier=e.quartier, type_bien=e.type_bien,
        surface=e.surface, pieces=e.pieces, etage=e.etage, annee=e.annee,
        etat=e.etat, balcon=e.balcon, balcon_pond=e.balcon_pond,
        parking_nb=e.parking_nb, parking_val=e.parking_val,
        prix_m2=e.prix_m2, marge=e.marge, description=e.description,
        atouts=e.atouts, inconvenients=e.inconvenients, courtier=e.courtier,
        notes=f"Copie de {e.id}"
    )
    db.session.add(new_e)
    db.session.commit()
    log_audit("cloned", new_e.id, f"Clonée de {e.id}")
    return redirect(url_for("estimation_report", eid=new_e.id))


@app.route("/estimation/<int:eid>/notes", methods=["POST"])
@login_required
def estimation_notes(eid):
    """Éditer les notes d'une estimation."""
    e = own_estimation_or_404(eid)
    e.notes = request.form.get("notes", "")
    db.session.commit()
    log_audit("updated_notes", e.id)
    flash("Notes mises à jour.")
    return redirect(url_for("estimation_report", eid=e.id))


APARTMENT_TYPES = {"Appartement", "Duplex", "Attique", "Penthouse", "Triplex"}
HOUSE_TYPES = {"Maison individuelle", "Villa", "Maison", "Hôtel Particulier"}


def _classify_type(type_bien):
    if not type_bien:
        return "autres"
    if type_bien in APARTMENT_TYPES:
        return "appartements"
    if type_bien in HOUSE_TYPES:
        return "maisons"
    return "autres"


@app.route("/classeur")
@login_required
def classeur():
    """Archive de toutes les estimations de l'utilisateur, groupée par quartier puis catégorie."""
    estimations = own_estimations().order_by(Estimation.created_at.desc()).all()
    grouped = {}  # quartier -> {"appartements": [], "maisons": [], "autres": []}
    for e in estimations:
        q = e.quartier or "Sans quartier"
        grouped.setdefault(q, {"appartements": [], "maisons": [], "autres": []})
        grouped[q][_classify_type(e.type_bien)].append(e)

    # Ordre voulu : les 3 quartiers principaux d'abord, puis le reste alphabétique
    priority = ["Champel", "Eaux-Vives", "Miremont"]
    ordered_quartiers = [q for q in priority if q in grouped]
    ordered_quartiers += sorted(q for q in grouped if q not in priority)

    return render_template(
        "classeur.html",
        grouped=grouped,
        quartiers=ordered_quartiers,
        total=len(estimations),
    )


@app.route("/dashboard")
@login_required
def dashboard():
    """Tableau de bord de l'utilisateur — ses estimations et un CTA nouvelle estimation."""
    uid = session.get("user_id")
    total = own_estimations().count()
    # prix_presentation est une @property Python : on approxime en SQL par prix_m2*surface*(1+marge)
    by_quartier = db.session.query(
        Estimation.quartier,
        db.func.count(Estimation.id),
        db.func.avg(Estimation.prix_m2),
        db.func.avg(Estimation.prix_m2 * Estimation.surface * (1 + Estimation.marge))
    ).filter(Estimation.user_id == uid).group_by(Estimation.quartier).all()

    recent = own_estimations().order_by(Estimation.created_at.desc()).limit(10).all()

    # Audit logs restreints aux estimations de l'utilisateur.
    own_ids = [row[0] for row in db.session.query(Estimation.id).filter_by(user_id=uid).all()]
    logs = (AuditLog.query.filter(AuditLog.estimation_id.in_(own_ids))
            .order_by(AuditLog.created_at.desc()).limit(20).all()) if own_ids else []

    return render_template("dashboard.html", total=total, by_quartier=by_quartier,
                         recent=recent, logs=logs, user=current_user())


@app.route("/estimation/<int:eid>/export.pdf")
@login_required
def estimation_export_pdf(eid):
    """Exporte une estimation en PDF (via HTML printable)."""
    e = own_estimation_or_404(eid)
    _, pool = ref_for(e.quartier, e.address)
    sold = [r for r in pool if r.kind in ("sold", "retenu")]
    forsale = [r for r in pool if r.kind == "forsale"]
    html = build_report(e, sold, forsale)

    return render_template("export_pdf.html", e=e, report_html=html)


@app.route("/estimation/<int:eid>/export.docx")
@login_required
def estimation_export_docx(eid):
    """Exporte une estimation en Word (.docx) au format LP — en-tête logo, sections, signature."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from html.parser import HTMLParser

    e = own_estimation_or_404(eid)

    # On s'assure d'avoir un corps de rapport (Claude ou fallback)
    _, comparables, stats = find_comparables(
        e.quartier, e.address, e.surface, e.type_bien,
        current_eid=e.id, user_id=session.get("user_id"),
    )
    body_html = _generate_lp_report(e, comparables, stats)
    if not body_html:
        _, legacy_pool = ref_for(e.quartier, e.address)
        sold_l = [r for r in legacy_pool if r.kind in ("sold", "retenu")]
        forsale_l = [r for r in legacy_pool if r.kind == "forsale"]
        body_html = build_report(e, sold_l, forsale_l)

    doc = Document()

    # Marges A4 raisonnables
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    # Style de base
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(11)

    LP_GOLD = RGBColor(0xB8, 0x89, 0x5B)
    LP_INK = RGBColor(0x2A, 0x25, 0x20)

    def _set_heading(paragraph, size=16, color=LP_INK, bold=True):
        for run in paragraph.runs:
            run.font.name = "Georgia"
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = color

    # ------ EN-TÊTE ------
    logo_path = os.path.join(app.config["UPLOAD_FOLDER"], DEFAULT_LOGO_FILENAME)
    if os.path.exists(logo_path):
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_p.add_run()
        try:
            run.add_picture(logo_path, width=Inches(1.4))
        except Exception:
            pass

    brand_p = doc.add_paragraph()
    brand_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brand_p.add_run("LEONARD PROPERTIES")
    run.font.name = "Georgia"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = LP_INK

    tag_p = doc.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tag_p.add_run("Rapport d'estimation hédoniste  ·  " + e.created_at.strftime("%d.%m.%Y"))
    run.font.name = "Georgia"
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = LP_GOLD

    # Titre bien
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(e.address or "Bien à estimer")
    run.font.name = "Georgia"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = LP_INK

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_parts = []
    if e.type_bien: subtitle_parts.append(e.type_bien)
    if e.surface: subtitle_parts.append(f"{int(e.surface)} m²")
    if e.quartier: subtitle_parts.append(e.quartier)
    run = subtitle_p.add_run("  ·  ".join(subtitle_parts))
    run.font.name = "Georgia"
    run.font.size = Pt(11)
    run.font.color.rgb = LP_GOLD

    # ------ SYNTHÈSE ------
    doc.add_paragraph()
    tbl = doc.add_table(rows=2, cols=3)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl.autofit = True
    headers = ["Valeur vénale", "Prix / m² retenu", "Prix de présentation"]
    values = [
        f"CHF {int(e.valeur_venale or 0):,}".replace(",", "'"),
        f"CHF {int(e.prix_m2 or 0):,}".replace(",", "'"),
        f"CHF {int(e.prix_presentation or 0):,}".replace(",", "'"),
    ]
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = "Georgia"; run.font.size = Pt(9); run.font.bold = True
        run.font.color.rgb = LP_GOLD
    for i, v in enumerate(values):
        cell = tbl.rows[1].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(v)
        run.font.name = "Georgia"; run.font.size = Pt(14); run.font.bold = True
        run.font.color.rgb = LP_INK

    # ------ COMPARABLES (max 3) ------
    if comparables:
        doc.add_paragraph()
        h = doc.add_paragraph("Comparables retenus")
        _set_heading(h, size=15, color=LP_INK)
        for c in comparables[:3]:
            p = doc.add_paragraph()
            r = p.add_run(c["adresse"])
            r.font.name = "Georgia"; r.font.size = Pt(12); r.font.bold = True
            r.font.color.rgb = LP_INK
            sub = []
            if c.get("type_bien"): sub.append(c["type_bien"])
            if c.get("annee"): sub.append(str(c["annee"]))
            if c.get("surface"): sub.append(f"{int(c['surface'])} m²")
            if sub:
                p2 = doc.add_paragraph()
                r = p2.add_run(" · ".join(sub))
                r.font.name = "Georgia"; r.font.size = Pt(10); r.font.italic = True
                r.font.color.rgb = LP_GOLD
            prix_line = []
            if c.get("prix_total"):
                prix_line.append(f"CHF {int(c['prix_total']):,}".replace(",", "'"))
            if c.get("prix_m2"):
                prix_line.append(f"{int(c['prix_m2']):,} CHF/m²".replace(",", "'"))
            if prix_line:
                p3 = doc.add_paragraph()
                r = p3.add_run(" · ".join(prix_line))
                r.font.name = "Georgia"; r.font.size = Pt(10); r.font.bold = True
                r.font.color.rgb = LP_INK
            if c.get("description"):
                p4 = doc.add_paragraph()
                r = p4.add_run(" ".join((c["description"]).split()))
                r.font.name = "Georgia"; r.font.size = Pt(10); r.font.italic = True

    # ------ CORPS DU RAPPORT (parse HTML minimaliste) ------
    doc.add_paragraph()

    class _HtmlToDocx(HTMLParser):
        def __init__(self, document):
            super().__init__()
            self.doc = document
            self.stack = []
            self.current_para = None
            self.in_ul = False
            self.bold = False
            self.italic = False

        def handle_starttag(self, tag, attrs):
            t = tag.lower()
            if t in ("h2",):
                self.current_para = self.doc.add_paragraph()
                self.stack.append(("h2", self.current_para))
            elif t in ("h3",):
                self.current_para = self.doc.add_paragraph()
                self.stack.append(("h3", self.current_para))
            elif t == "p":
                self.current_para = self.doc.add_paragraph()
                self.stack.append(("p", self.current_para))
            elif t == "ul":
                self.in_ul = True
            elif t == "li":
                self.current_para = self.doc.add_paragraph(style="List Bullet")
                self.stack.append(("li", self.current_para))
            elif t in ("strong", "b"):
                self.bold = True
            elif t in ("em", "i"):
                self.italic = True

        def handle_endtag(self, tag):
            t = tag.lower()
            if t in ("h2", "h3", "p", "li") and self.stack:
                kind, para = self.stack.pop()
                if kind == "h2":
                    _set_heading(para, size=15, color=LP_INK)
                elif kind == "h3":
                    _set_heading(para, size=12, color=LP_GOLD)
                self.current_para = None
            elif t == "ul":
                self.in_ul = False
            elif t in ("strong", "b"):
                self.bold = False
            elif t in ("em", "i"):
                self.italic = False

        def handle_data(self, data):
            if not data.strip() and self.current_para is None:
                return
            if self.current_para is None:
                self.current_para = self.doc.add_paragraph()
            run = self.current_para.add_run(data)
            run.font.name = "Georgia"
            run.font.size = Pt(11)
            run.bold = self.bold
            run.italic = self.italic

    parser = _HtmlToDocx(doc)
    parser.feed(body_html)

    # ------ ANNEXE : PANEL COMPLET DE COMPARABLES ------
    if comparables:
        doc.add_paragraph()
        h = doc.add_paragraph(f"Panel de biens comparables ({len(comparables)})")
        _set_heading(h, size=15, color=LP_INK)

        ptbl = doc.add_table(rows=1, cols=6)
        ptbl.autofit = True
        headers = ["Adresse", "Type", "Année", "Surface", "Prix total", "Prix / m²"]
        for i, htxt in enumerate(headers):
            cell = ptbl.rows[0].cells[i]
            p = cell.paragraphs[0]
            run = p.add_run(htxt)
            run.font.name = "Georgia"; run.font.size = Pt(9); run.font.bold = True
            run.font.color.rgb = LP_GOLD
        for c in comparables:
            row = ptbl.add_row().cells
            def _fill(cell, txt, bold=False, color=LP_INK):
                p = cell.paragraphs[0]
                r = p.add_run(txt)
                r.font.name = "Georgia"; r.font.size = Pt(9); r.font.bold = bold
                r.font.color.rgb = color
            _fill(row[0], c.get("adresse") or "—", bold=True)
            _fill(row[1], c.get("type_bien") or "—")
            _fill(row[2], str(c.get("annee") or "—"))
            _fill(row[3], f"{int(c['surface'])} m²" if c.get("surface") else "—")
            _fill(row[4], f"CHF {int(c['prix_total']):,}".replace(",", "'") if c.get("prix_total") else "—")
            _fill(row[5], f"{int(c['prix_m2']):,} CHF".replace(",", "'") if c.get("prix_m2") else "—",
                  bold=True, color=LP_GOLD)

    # ------ SIGNATURE ------
    doc.add_paragraph()
    sig_p = doc.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sig_p.add_run(e.courtier or "LEONARD PROPERTIES SA")
    run.font.name = "Georgia"; run.font.size = Pt(11); run.font.bold = True
    run.font.color.rgb = LP_INK

    # Export
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"estimation-{(e.address or 'bien').lower().replace(' ', '-')[:50]}-{e.id}.docx"
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


@app.route("/import-csv", methods=["GET", "POST"])
@login_required
def import_csv():
    """Importe des estimations depuis un fichier CSV."""
    if request.method == "POST":
        if "file" not in request.files:
            flash("Aucun fichier sélectionné.")
            return redirect(url_for("import_csv"))

        file = request.files["file"]
        if not file.filename.endswith(".csv"):
            flash("Veuillez uploader un fichier CSV.")
            return redirect(url_for("import_csv"))

        try:
            stream = io.StringIO(file.read().decode("utf-8"))
            reader = csv.DictReader(stream)

            count = 0
            for row in reader:
                def num(k, d=0.0):
                    try:
                        v = row.get(k, "")
                        return float(str(v).replace("'", "").replace(",", ".") or d)
                    except (ValueError, TypeError):
                        return d

                est = Estimation(
                    address=row.get("address", ""),
                    quartier=row.get("quartier", ""),
                    type_bien=row.get("type_bien", ""),
                    surface=num("surface"),
                    pieces=row.get("pieces", ""),
                    etage=row.get("etage", ""),
                    annee=row.get("annee", ""),
                    etat=row.get("etat", ""),
                    balcon=num("balcon"),
                    balcon_pond=num("balcon_pond", 0.5),
                    parking_nb=int(num("parking_nb")),
                    parking_val=num("parking_val"),
                    prix_m2=num("prix_m2"),
                    marge=num("marge", 0.07),
                    description=row.get("description", ""),
                    atouts=row.get("atouts", ""),
                    inconvenients=row.get("inconvenients", ""),
                    courtier=row.get("courtier", ""),
                    notes=row.get("notes", "")
                )
                db.session.add(est)
                count += 1

            db.session.commit()
            flash(f"{count} estimations importées avec succès!")
            return redirect(url_for("estimations_list"))
        except Exception as e:
            flash(f"Erreur lors de l'import: {str(e)}")
            return redirect(url_for("import_csv"))

    return render_template("import_csv.html")


@app.route("/theme/<theme>")
@login_required
def set_theme(theme):
    """Définit le thème (dark/light)."""
    if theme in ["dark", "light"]:
        session["theme"] = theme
    return redirect(request.referrer or url_for("index"))


@app.route("/estimation/<int:eid>/export.csv")
@login_required
def estimation_export_csv(eid):
    """Exporte une estimation en CSV."""
    e = own_estimation_or_404(eid)

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Propriété", "Valeur"])
    writer.writerow(["Adresse", e.address])
    writer.writerow(["Quartier", e.quartier])
    writer.writerow(["Type de bien", e.type_bien])
    writer.writerow(["Surface", e.surface])
    writer.writerow(["Pièces", e.pieces])
    writer.writerow(["État", e.etat])
    writer.writerow(["Balcon", e.balcon])
    writer.writerow(["Parking", e.parking_nb])
    writer.writerow(["Prix/m²", e.prix_m2])
    writer.writerow(["Valeur estimée", e.prix_presentation])
    writer.writerow(["Notes", e.notes])

    output.seek(0)
    return send_file(
        io.BytesIO(output.getvalue().encode("utf-8")),
        mimetype="text/csv",
        as_attachment=True,
        download_name=f"estimation_{e.id}.csv"
    )


@app.route("/api/backup")
@login_required
def api_backup():
    """Backup complet de la base — users, estimations, refs, settings, audit logs.
    Retourne un JSON téléchargeable, à archiver hors-Railway pour parer à une perte de conteneur.
    """
    def _serialize_est(e):
        return {
            "id": e.id, "user_id": e.user_id,
            "created_at": e.created_at.isoformat() if e.created_at else None,
            "address": e.address, "quartier": e.quartier, "type_bien": e.type_bien,
            "surface": e.surface, "pieces": e.pieces, "etage": e.etage,
            "annee": e.annee, "etat": e.etat,
            "balcon": e.balcon, "balcon_pond": e.balcon_pond,
            "parking_nb": e.parking_nb, "parking_val": e.parking_val,
            "prix_m2": e.prix_m2, "marge": e.marge,
            "description": e.description, "atouts": e.atouts,
            "inconvenients": e.inconvenients, "courtier": e.courtier,
            "notes": e.notes,
            "valeur_venale": e.valeur_venale,
            "prix_presentation": e.prix_presentation,
        }
    payload = {
        "meta": {
            "generated_at": datetime.utcnow().isoformat() + "Z",
            "app": "lp-estimation",
            "version": "1.0",
        },
        "users": [
            {"id": u.id, "email": u.email, "name": u.name,
             "created_at": u.created_at.isoformat() if u.created_at else None}
            for u in User.query.all()
        ],
        "estimations": [_serialize_est(e) for e in Estimation.query.all()],
        "ref_prices": [
            {"id": r.id, "quartier": r.quartier, "adresse": r.adresse,
             "type_bien": r.type_bien, "surface": r.surface,
             "prix_m2": r.prix_m2, "prix_total": r.prix_total,
             "annee": r.annee, "source": r.source, "kind": r.kind,
             "description": r.description, "reference": r.reference}
            for r in RefPrice.query.all()
        ],
        "settings": [{"key": s.key, "value": s.value} for s in Setting.query.all()],
        "audit_logs": [
            {"id": l.id, "action": l.action, "estimation_id": l.estimation_id,
             "created_at": l.created_at.isoformat() if l.created_at else None,
             "description": l.description}
            for l in AuditLog.query.all()
        ],
    }
    import json
    filename = f"backup_lp-estimation_{datetime.utcnow().strftime('%Y%m%d_%H%M')}.json"
    return send_file(
        io.BytesIO(json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")),
        mimetype="application/json",
        as_attachment=True,
        download_name=filename,
    )


@app.route("/export-all.csv")
@login_required
def export_all_csv():
    """Exporte toutes les estimations de l'utilisateur en CSV."""
    estimations = own_estimations().all()

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([
        "address", "quartier", "type_bien", "surface", "pieces", "etage", "annee",
        "etat", "balcon", "balcon_pond", "parking_nb", "parking_val", "prix_m2",
        "marge", "description", "atouts", "inconvenients", "courtier", "notes"
    ])

    for e in estimations:
        writer.writerow([
            e.address, e.quartier, e.type_bien, e.surface, e.pieces, e.etage, e.annee,
            e.etat, e.balcon, e.balcon_pond, e.parking_nb, e.parking_val, e.prix_m2,
            e.marge, e.description, e.atouts, e.inconvenients, e.courtier, e.notes
        ])

    output.seek(0)
    return send_file(
        io.BytesIO(output.getvalue().encode("utf-8")),
        mimetype="text/csv",
        as_attachment=True,
        download_name=f"estimations_all.csv"
    )


# ----------------------- FILTRES ET CONTEXTE -----------------------
@app.context_processor
def inject_helpers():
    """Injecte les fonctions helper dans les templates."""
    return dict(get_logo_path=get_logo_path)


@app.template_filter("chf")
def chf(v):
    try:
        return "{:,.0f}".format(float(v)).replace(",", "'")
    except (ValueError, TypeError):
        return "—"


@app.template_filter("m2")
def m2(v):
    try:
        s = "{:,.1f}".format(float(v)).replace(",", "'")
        return s.rstrip("0").rstrip(".")
    except (ValueError, TypeError):
        return "—"


CSV_SOURCE_TAG = "CSV Belfin/Champel60"
CSV_LP_SOURCE_TAG = "CSV LP consolidé"
CSV_ENRICHED_TAG = "CSV enrichi (Champel60)"
CSV_ESTIMATIONS_LP_TAG = "Estimation LP passée"

CSV_FILES = [
    # (chemin_relatif, source_tag, format)
    #   format: 'basic' | 'enriched' | 'estimation_lp'
    ("data/comparables.csv", CSV_SOURCE_TAG, "basic"),
    ("data/comparables-lp.csv", CSV_LP_SOURCE_TAG, "basic"),
    ("data/comparables_enrichis.csv", CSV_ENRICHED_TAG, "enriched"),
    ("data/comparables_estimations_lp.csv", CSV_ESTIMATIONS_LP_TAG, "estimation_lp"),
]


def _normalize_type(t):
    """Nettoie les variations 'Villa Individuelle' / 'villa individuelle' / etc.
    Retourne la version canonique ou t si inconnu."""
    if not t: return ""
    tt = t.strip().lower()
    if "hôtel particulier" in tt or "hp" == tt: return "Hôtel Particulier"
    if "château" in tt: return "Villa"
    if "villa" in tt: return "Villa"
    if "maison" in tt: return "Maison individuelle"
    if "townhouse" in tt: return "Maison individuelle"
    if "attique" in tt: return "Attique"
    if "penthouse" in tt: return "Penthouse"
    if "duplex" in tt: return "Duplex"
    if "triplex" in tt: return "Triplex"
    if "appartement" in tt or "ppe" in tt: return "Appartement"
    if "parcelle" in tt or "terrain" in tt: return ""  # skip non-résidentiel
    return t.strip()


def _parse_swiss_number(v):
    """Parse '6’100’000' ou '13'555' → 6100000. Retourne None si invalide."""
    if v is None:
        return None
    s = str(v).strip()
    if not s:
        return None
    # Enlève apostrophes typographiques ’ ' et espaces
    s = s.replace("’", "").replace("'", "").replace(" ", "").replace(" ", "")
    try:
        return float(s)
    except ValueError:
        return None


def load_comparables_csv():
    """Charge (ou recharge) tous les CSV comparables listés dans CSV_FILES.

    Idempotent : supprime d'abord toutes les entrées taggées avec un de nos tags,
    puis réinjecte. Ainsi, modifier les CSV et redéployer suffit à mettre à jour la base.
    """
    import csv as _csv
    base = os.path.dirname(__file__)

    # Purge des anciennes entrées CSV (peu importe le tag)
    RefPrice.query.filter(RefPrice.source.in_([tag for _, tag, _ in CSV_FILES])).delete(synchronize_session=False)
    db.session.commit()

    total = 0
    for rel_path, source_tag, fmt in CSV_FILES:
        path = os.path.join(base, rel_path)
        if not os.path.exists(path):
            app.logger.info(f"CSV absent ({path}), skip.")
            continue
        inserted = 0
        with open(path, encoding="utf-8") as f:
            for row in _csv.DictReader(f):
                try:
                    if fmt == "estimation_lp":
                        # Format : address, quartier, type_bien, surface, annee,
                        #          prix_m2_retenu, valeur_venale, prix_presentation, description
                        adresse = " ".join((row.get("address") or "").split()).strip()[:200]
                        if not adresse:
                            continue
                        type_bien = _normalize_type(row.get("type_bien") or "")[:80]
                        if not type_bien:
                            continue  # non résidentiel, skip
                        surface = _parse_swiss_number(row.get("surface"))
                        prix_m2 = _parse_swiss_number(row.get("prix_m2_retenu"))
                        valeur = _parse_swiss_number(row.get("valeur_venale"))
                        pres = _parse_swiss_number(row.get("prix_presentation"))
                        # Recalcule prix_m2 si absent mais on a valeur/surface
                        if not prix_m2 and valeur and surface:
                            prix_m2 = round(valeur / surface)
                        if not prix_m2 and pres and surface:
                            # prix de présentation = valeur × (1 + marge ~7 %)
                            prix_m2 = round(pres / 1.07 / surface)
                        if not prix_m2:
                            continue
                        # Extrait une année propre (4 chiffres, la plus récente si plusieurs).
                        # La colonne annee = VARCHAR(20), certains rapports contiennent
                        # 'Construction 2011-2015, rénové en 2024' → trop long, on prend '2024'.
                        import re as _re
                        annee_str = (row.get("annee") or "").strip()
                        yrs = _re.findall(r"(19\d\d|20\d\d)", annee_str)
                        annee = max(yrs) if yrs else annee_str[:20]
                        r = RefPrice(
                            quartier=(row.get("quartier") or "").strip()[:120],
                            adresse=adresse,
                            prix_m2=prix_m2,
                            prix_total=valeur or pres,
                            annee=annee,
                            source=source_tag,
                            kind="retenu",  # estimation LP passée = prix retenu par le courtier
                            surface=surface,
                            type_bien=type_bien,
                            description=(row.get("description") or "").strip() or None,
                            reference=None,
                        )
                    elif fmt == "enriched":
                        # Format enrichi : surface_m2, prix_chf, prix_m2 (avec apostrophes),
                        # description multi-lignes, reference
                        surface = _parse_swiss_number(row.get("surface_m2"))
                        prix_total = _parse_swiss_number(row.get("prix_chf"))
                        prix_m2 = _parse_swiss_number(row.get("prix_m2"))
                        # Calculer prix/m² si absent
                        if not prix_m2 and prix_total and surface:
                            prix_m2 = round(prix_total / surface)
                        if not prix_m2:
                            continue
                        # Nettoyer l'adresse (peut contenir des retours à la ligne)
                        adresse = " ".join((row.get("adresse") or "").split())
                        r = RefPrice(
                            quartier=(row.get("quartier") or "").strip(),
                            adresse=adresse,
                            prix_m2=prix_m2,
                            prix_total=prix_total,
                            annee=(row.get("annee") or "").strip(),
                            source=source_tag,
                            kind="sold",
                            surface=surface,
                            type_bien=(row.get("type_bien") or "").strip() or None,
                            description=(row.get("description") or "").strip() or None,
                            reference=(row.get("reference") or "").strip() or None,
                        )
                    else:  # basic
                        prix_m2 = float(row.get("prix_m2") or 0)
                        surface = float(row.get("surface")) if row.get("surface") else None
                        if not prix_m2:
                            continue
                        r = RefPrice(
                            quartier=(row.get("quartier") or "").strip(),
                            adresse=(row.get("adresse") or "").strip(),
                            prix_m2=prix_m2,
                            annee=(row.get("annee") or "").strip(),
                            source=source_tag,
                            kind="sold",
                            surface=surface,
                            type_bien=(row.get("type_bien") or "").strip() or None,
                        )
                    db.session.add(r)
                    inserted += 1
                except Exception as e:
                    app.logger.warning(f"Ligne CSV ignorée : {row} — {e}")
        db.session.commit()
        app.logger.info(f"CSV comparables : {inserted} lignes chargées depuis {rel_path}")
        total += inserted
    return total


def seed():
    db.create_all()
    if RefPrice.query.first():
        return
    # Références prix/m² APPARTEMENTS, extraites des estimations LP fournies
    data = [
        # quartier, adresse, prix_m2, annee, source, kind
        ("Pâquis", "Abraham-Gevray 1 (lots 3.03/4.03)", 15527, "2025", "Estimation Gevray 1", "sold"),
        ("Pâquis", "Abraham-Gevray 1 (lots 5.05/6.04)", 19567, "2024", "Estimation Gevray 1", "sold"),
        ("Pâquis", "Abraham-Gevray 1 (lot 5.07)", 19215, "2024", "Estimation Gevray 1", "sold"),
        ("Pâquis", "Abraham-Gevray 1 (attique 10.01/9.01)", 29319, "2024", "Estimation Gevray 1", "sold"),
        ("Pâquis", "Abraham-Gevray 1 (lots 5.02/6.02)", 21343, "2022", "Estimation Gevray 1", "sold"),
        ("Pâquis", "Abraham-Gevray 1 (retenu)", 22000, "2026", "Estimation Gevray 1 — prix retenu", "retenu"),
        ("Eaux-Vives", "Rue Abraham-Constantin 4-6", 19494, "2024", "Estimation Florissant 47", "sold"),
        ("Eaux-Vives", "Avenue Alfred-Bertrand 13", 21531, "2024", "Estimation Florissant 47", "sold"),
        ("Eaux-Vives", "Avenue Peschier 24", 15918, "2023", "Estimation Florissant 47", "sold"),
        ("Eaux-Vives", "Route de Florissant 47 (retenu)", 19876, "2025", "Estimation Florissant 47 — prix retenu", "retenu"),
        ("Champel", "Avenue de Champel 14", 21120, "2023", "Estimation Florissant 47", "sold"),
        ("Champel", "Avenue de Champel 14", 22719, "2022", "Estimation Florissant 47", "sold"),
        ("Champel", "Rue Monnier 1", 18478, "2023", "Estimation Florissant 47", "forsale"),
        ("Champel", "Chemin Tour de Champel 12", 17883, "2025", "Estimation Champel 60", "forsale"),
        ("Champel", "Avenue de Miremont 30 (retenu)", 14000, "2026", "Estimation Miremont 30 — prix retenu", "retenu"),
    ]
    for q, a, pm, an, src, k in data:
        db.session.add(RefPrice(quartier=q, adresse=a, prix_m2=pm, annee=an, source=src, kind=k))
    db.session.commit()


def _migrate():
    """Ajoute les colonnes ajoutées après le déploiement initial (sans framework de migration)."""
    from sqlalchemy import text
    for stmt in (
        "ALTER TABLE estimation ADD COLUMN user_id INTEGER",
        "ALTER TABLE estimation ADD COLUMN report_ai TEXT",
        "ALTER TABLE ref_price ADD COLUMN surface FLOAT",
        "ALTER TABLE ref_price ADD COLUMN type_bien VARCHAR(80)",
        "ALTER TABLE ref_price ADD COLUMN prix_total FLOAT",
        "ALTER TABLE ref_price ADD COLUMN description TEXT",
        "ALTER TABLE ref_price ADD COLUMN reference VARCHAR(80)",
    ):
        try:
            db.session.execute(text(stmt))
            db.session.commit()
        except Exception:
            db.session.rollback()


REPORT_PROMPT_VERSION = "v3-lp-sonnet5-2026-07-03"


def _invalidate_stale_ai_reports():
    """Vide report_ai pour toutes les estimations quand le prompt LP a changé de version.
    Idempotent : ne s'exécute qu'une fois par version, grâce à la table Setting."""
    key = "report_prompt_version"
    setting = Setting.query.filter_by(key=key).first()
    if setting and setting.value == REPORT_PROMPT_VERSION:
        return
    # Prompt changé → purge les rapports IA existants pour forcer une régénération
    Estimation.query.update({Estimation.report_ai: None}, synchronize_session=False)
    if setting:
        setting.value = REPORT_PROMPT_VERSION
    else:
        db.session.add(Setting(key=key, value=REPORT_PROMPT_VERSION))
    db.session.commit()
    app.logger.info(f"Rapports IA invalidés (nouveau prompt {REPORT_PROMPT_VERSION})")


with app.app_context():
    db.create_all()
    _migrate()
    seed()
    # non bloquants : si l'un plante, on log et on continue à démarrer
    try:
        load_comparables_csv()
    except Exception as _e:
        app.logger.exception(f"load_comparables_csv KO au démarrage : {_e}")
    try:
        _invalidate_stale_ai_reports()
    except Exception as _e:
        app.logger.exception(f"_invalidate_stale_ai_reports KO : {_e}")


if __name__ == "__main__":
    app.run(debug=True, port=5002)
